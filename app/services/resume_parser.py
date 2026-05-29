import io
import re
import fitz  # PyMuPDF
import docx

# ─────────────────────────────────────────────
# Logger (same as groq_service for consistency)
# ─────────────────────────────────────────────
def log(tag: str, message: str, level: str = "INFO"):
    levels = {
        "INFO":   "\033[94m[INFO]\033[0m",
        "OK":     "\033[92m[OK]\033[0m",
        "WARN":   "\033[93m[WARN]\033[0m",
        "ERROR":  "\033[91m[ERROR]\033[0m",
        "RESUME": "\033[96m[RESUME]\033[0m",
    }
    prefix = levels.get(level, "[LOG]")
    print(f"{prefix} [{tag}] {message}", flush=True)


class ResumeParser:

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        log("RESUME", f"Parsing PDF — {len(file_bytes)} bytes received", "RESUME")
        text = ""
        page_count = 0

        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                page_count = len(doc)
                log("RESUME", f"PDF opened successfully — {page_count} page(s) found", "OK")

                for i, page in enumerate(doc):
                    page_text = page.get_text()
                    word_count = len(page_text.split())
                    log("RESUME", f"  Page {i+1}: extracted ~{word_count} words", "INFO")
                    text += page_text + "\n"

        except Exception as e:
            log("RESUME", f"❌ Failed to parse PDF: {e}", "ERROR")
            return ""

        cleaned = ResumeParser.clean_text(text)
        final_word_count = len(cleaned.split())

        if final_word_count < 30:
            log("RESUME", f"⚠️  Very little text extracted ({final_word_count} words) — PDF may be image-based/scanned!", "WARN")
            log("RESUME", "Consider adding OCR support (e.g. pytesseract) for scanned PDFs", "WARN")
        else:
            log("RESUME", f"✅ PDF parsed successfully — {final_word_count} words extracted across {page_count} page(s)", "OK")

        log("RESUME", f"Text preview (first 300 chars): {cleaned[:300]}", "INFO")
        return cleaned

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        log("RESUME", f"Parsing DOCX — {len(file_bytes)} bytes received", "RESUME")
        text = ""
        para_count = 0

        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            para_count = len(doc.paragraphs)
            log("RESUME", f"DOCX opened — {para_count} paragraph(s) found", "OK")

            for para in doc.paragraphs:
                if para.text.strip():  # Skip blank paragraphs
                    text += para.text + "\n"

        except Exception as e:
            log("RESUME", f"❌ Failed to parse DOCX: {e}", "ERROR")
            return ""

        cleaned = ResumeParser.clean_text(text)
        final_word_count = len(cleaned.split())

        if final_word_count < 30:
            log("RESUME", f"⚠️  Very little text extracted ({final_word_count} words) — DOCX may be corrupted or mostly images!", "WARN")
        else:
            log("RESUME", f"✅ DOCX parsed successfully — {final_word_count} words from {para_count} paragraphs", "OK")

        log("RESUME", f"Text preview (first 300 chars): {cleaned[:300]}", "INFO")
        return cleaned

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Cleans resume text while preserving meaningful structure.

        Old approach collapsed ALL whitespace into single spaces, which destroyed
        bullet points and section breaks — making the resume hard for the LLM to parse.

        New approach:
        - Normalizes multiple spaces/tabs → single space
        - Preserves single newlines (section separators, bullet points)
        - Collapses 3+ newlines → double newline (section gaps)
        - Removes non-printable characters
        """
        if not text:
            return ""

        # Remove non-printable characters (except newlines and tabs)
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')

        # Normalize tabs and multiple spaces to single space (within a line)
        text = re.sub(r'[ \t]+', ' ', text)

        # Collapse 3+ consecutive newlines to 2 (preserve section breaks)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]

        # Remove lines that are pure whitespace
        lines = [line for line in lines if line]

        result = '\n'.join(lines).strip()

        log("RESUME", f"Text cleaned — {len(result)} chars remaining", "INFO")
        return result


resume_parser = ResumeParser()